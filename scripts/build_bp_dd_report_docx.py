#!/usr/bin/env python3
"""
BP DD Report DOCX 生成器 v3

v3 改进：
- 来源列表从各维度正文中剥离，统一编号放到报告末尾
- 表格美化：表头底纹、列宽自适应、单元格内边距、紧凑字号
- 完整 Markdown→DOCX 转换（表格、行内格式、引用块、有序列表、分隔线）
- 来源引注清洗（移除内部文件路径引用）
- 封面不暴露 task ID
"""
from __future__ import annotations

import json
import re
import time
from pathlib import Path
from typing import Any


# ── 东亚字体设置 ──────────────────────────────────────

def _set_eastasia_font_on_style(style, font_name: str):
    """在样式级别设置 eastAsia 字体（解决 MS 明朝问题）"""
    from docx.oxml.ns import qn
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _set_eastasia_font_on_run(run, font_name: str):
    """在 run 级别设置 eastAsia 字体"""
    from docx.oxml.ns import qn
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


# ── 内部信息清洗 ──────────────────────────────────────

_INTERNAL_SOURCE_PATTERNS = [
    r"bp_ocr_text\.txt",
    r"bp_step0_profile\.(json|md)",
    r"company_verify_report\.(json|md)",
    r"bp_presearch_step_\w+\.md",
    r"bp_presearch_results\.json",
    r"bp_phase\d+_\w+\.(md|json)",
    r"phase\d+_dispatch\.json",
    r"bp_phase2_brief_\w+\.md",
    r"bp_phase2_manifest_\w+\.json",
    r"bp_phase2_spawn_\w+\.json",
]


def _is_internal_source_row(row_text: str) -> bool:
    for pat in _INTERNAL_SOURCE_PATTERNS:
        if re.search(pat, row_text, re.IGNORECASE):
            return True
    return False


def _sanitize_text(text: str) -> str:
    return re.sub(r"TASK-\d{8}-\d{3}", "", text)


# ── 来源剥离 ──────────────────────────────────────────

def _strip_source_section(markdown: str) -> tuple[str, list[dict]]:
    """从 Markdown 中剥离"来源列表"章节，返回 (cleaned_text, sources)。

    sources 格式: [{"id": "S-1", "name": "xxx", "url": "xxx", "usage": "xxx"}, ...]
    
    同时提取 [^N] 格式的脚注定义（如 [^1]: 来源名 — URL），保留正文中的 [^N] 引用标记。
    """
    lines = markdown.split("\n")
    cleaned_lines = []
    sources: list[dict] = []
    in_source_section = False
    in_source_table = False
    # 也剥离末尾的免责声明（子代理自己写的，报告统一放）
    in_trailing_disclaimer = False
    
    # 第一遍：提取 [^N] 脚注定义
    footnote_defs: dict[str, dict] = {}
    remaining_lines = []
    in_footnote_defs = False
    
    for line in lines:
        stripped = line.strip()
        # 匹配 [^N]: xxx 格式的脚注定义
        fn_match = re.match(r"^\[\^(\d+)\]:\s*(.+)$", stripped)
        if fn_match:
            fn_id = fn_match.group(1)
            fn_content = fn_match.group(2).strip()
            # 尝试从脚注内容中提取 URL
            url_match = re.search(r"(https?://[^\s\)\]\"']+)", fn_content)
            url = url_match.group(1) if url_match else ""
            name = re.sub(r"\s*https?://[^\s\)\]\"']+\s*", "", fn_content).strip().rstrip("—–- ")
            footnote_defs[fn_id] = {
                "id": fn_id,
                "name": name or fn_content[:80],
                "url": url,
                "usage": "",
                "is_footnote": True,
            }
            continue
        remaining_lines.append(line)
    
    # 将脚注定义添加到 sources
    for fn_id in sorted(footnote_defs.keys(), key=int):
        sources.append(footnote_defs[fn_id])
    
    # 第二遍：处理来源列表章节（表格格式的来源）
    for line in remaining_lines:
        stripped = line.strip()

        # 检测来源列表章节开始
        if re.match(r"^#{1,3}\s+\d*\.?\s*来源", stripped):
            in_source_section = True
            in_source_table = False
            continue
        
        # 也检测"来源与参考"标题
        if re.match(r"^#{1,3}\s+\d*\.?\s*来源与参考", stripped):
            in_source_section = True
            in_source_table = False
            continue

        # 检测末尾免责声明
        if stripped.startswith("> 免责声明") or stripped.startswith(">免责声明"):
            in_trailing_disclaimer = True
            continue
        if in_trailing_disclaimer:
            if stripped.startswith(">") or not stripped:
                continue
            else:
                in_trailing_disclaimer = False

        if in_source_section:
            # 来源表格行
            if stripped.startswith("|"):
                # 跳过分隔行
                if re.match(r"^\|[\s\-:]+\|", stripped):
                    in_source_table = True
                    continue
                cells = [c.strip() for c in stripped.split("|")]
                if cells and cells[0] == "":
                    cells = cells[1:]
                if cells and cells[-1] == "":
                    cells = cells[:-1]

                if not in_source_table:
                    # 表头行，标记开始
                    in_source_table = True
                    continue

                # 数据行 — 提取来源
                if len(cells) >= 3:
                    row_text = " ".join(cells)
                    if _is_internal_source_row(row_text):
                        continue
                    src = {
                        "id": cells[0] if len(cells) > 0 else "",
                        "name": cells[1] if len(cells) > 1 else "",
                        "url": cells[2] if len(cells) > 2 else "",
                        "usage": cells[3] if len(cells) > 3 else "",
                    }
                    # 只保留有外部 URL 的来源
                    if src["url"].startswith("http"):
                        sources.append(src)
                    
                    # 也接受 [^N] 格式的来源（无 URL 但有名称）
                    if not src["url"].startswith("http") and src["name"] and re.match(r"^\[\^?\d+\]$", src["id"]):
                        sources.append(src)
                continue

            # [^N] 格式的来源列表（非表格，逐行列出的）
            fn_list_match = re.match(r"^\[\^(\d+)\]\s+(.+)$", stripped)
            if fn_list_match:
                fn_id = fn_list_match.group(1)
                fn_content = fn_list_match.group(2).strip()
                url_match = re.search(r"(https?://[^\s\)\]\"']+)", fn_content)
                url = url_match.group(1) if url_match else ""
                name = re.sub(r"\s*https?://[^\s\)\]\"']+\s*", "", fn_content).strip().rstrip("—–- ")
                # 避免重复添加
                if fn_id not in footnote_defs:
                    sources.append({
                        "id": fn_id,
                        "name": name or fn_content[:80],
                        "url": url,
                        "usage": "",
                        "is_footnote": True,
                    })
                continue

            # 来源章节内的分隔线或空行
            if not stripped or re.match(r"^-{3,}$", stripped):
                continue

            # 遇到新的章节标题，来源章节结束
            if stripped.startswith("#"):
                in_source_section = False
                cleaned_lines.append(line)
                continue

            # 来源章节内的其他内容（跳过）
            continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines), sources


# ── Markdown 解析工具 ─────────────────────────────────

def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|[\s\-:]+\|", line):
            i += 1
            continue
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        rows.append(cells)
        i += 1
    return rows, i


def _add_inline_formatted_text(paragraph, text: str):
    """解析行内 Markdown 格式并添加到段落。支持 [^N] 脚注标记渲染为上标。"""
    from docx.shared import Pt, RGBColor
    # 扩展 pattern 以匹配 [^N] 脚注标记
    pattern = r"(\*\*(.+?)\*\*|\[([^\]]+)\]\(([^)]+)\)|\[\^(\d+)\]|`([^`]+)`)"
    last_end = 0
    for m in re.finditer(pattern, text):
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            _set_eastasia_font_on_run(run, "宋体")
        if m.group(2) is not None:
            # **bold**
            run = paragraph.add_run(m.group(2))
            run.font.bold = True
            _set_eastasia_font_on_run(run, "宋体")
        elif m.group(3) is not None and m.group(5) is None:
            # [text](url) hyperlink
            run = paragraph.add_run(m.group(3))
            run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
            run.font.underline = True
            _set_eastasia_font_on_run(run, "宋体")
        elif m.group(5) is not None:
            # [^N] footnote reference → render as superscript
            fn_num = m.group(5)
            run = paragraph.add_run(f"[{fn_num}]")
            run.font.superscript = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
            _set_eastasia_font_on_run(run, "宋体")
        elif m.group(6) is not None:
            # `code`
            run = paragraph.add_run(m.group(6))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        last_end = m.end()
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        _set_eastasia_font_on_run(run, "宋体")


# ── 表格美化 ─────────────────────────────────────────

def _add_table_to_doc(doc, rows: list[list[str]]):
    """添加美化表格：表头底纹、紧凑字号、自适应列宽。"""
    from docx.shared import Pt, Cm, RGBColor, Inches, Emu
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 自动布局（让 Word 自适应列宽）
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    tblLayout = tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "autofit"})
    # 移除已有的 tblLayout
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    tblPr.append(tblLayout)

    # 表格总宽度 100%
    tblW = tblPr.makeelement(qn("w:tblW"), {qn("w:w"): "5000", qn("w:type"): "pct"})
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblPr.append(tblW)

    # 单元格默认内边距
    tblCellMar = tblPr.makeelement(qn("w:tblCellMar"), {})
    for side, val in [("top", "40"), ("bottom", "40"), ("left", "80"), ("right", "80")]:
        el = tblCellMar.makeelement(qn(f"w:{side}"), {qn("w:w"): val, qn("w:type"): "dxa"})
        tblCellMar.append(el)
    for existing in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(existing)
    tblPr.append(tblCellMar)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            _add_inline_formatted_text(p, cell_text.strip())

            # 字号
            for run in p.runs:
                run.font.size = Pt(9)
                _set_eastasia_font_on_run(run, "宋体")

            # 表头样式
            if row_idx == 0:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="1F4E79" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                if row_idx % 2 == 0:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F2F7FB" w:val="clear"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading)

    # 表格后间距
    p = doc.add_paragraph("")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# ── Markdown → DOCX 渲染 ─────────────────────────────

def _render_markdown_to_doc(doc, markdown_text: str):
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn

    lines = markdown_text.split("\n")
    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 分隔线
        if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "4",
                qn("w:space"): "1", qn("w:color"): "CCCCCC",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = _sanitize_text(heading_match.group(2).strip())
            doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        # 引用块
        if stripped.startswith("> "):
            quote_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            _add_inline_formatted_text(p, quote_text)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.italic = True
                run.font.size = Pt(10)
                _set_eastasia_font_on_run(run, "宋体")
            i += 1
            continue

        # 表格
        if stripped.startswith("|") and "|" in stripped[1:]:
            rows, end_i = _parse_table(lines, i)
            if rows and len(rows) > 1:
                _add_table_to_doc(doc, rows)
            i = end_i
            continue

        # 无序列表
        if stripped.startswith("- ") or stripped.startswith("* "):
            item_text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatted_text(p, item_text)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if ol_match:
            item_text = ol_match.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatted_text(p, item_text)
            i += 1
            continue

        # 纯加粗行
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.font.bold = True
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_inline_formatted_text(p, _sanitize_text(stripped))
        i += 1


# ── 来源清洗 ─────────────────────────────────────────

def _clean_dimension_output(text: str) -> str:
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and _is_internal_source_row(stripped):
            continue
        if stripped.startswith(">") and any(
            re.search(pat, stripped) for pat in _INTERNAL_SOURCE_PATTERNS
        ):
            cleaned.append("> 数据来源：BP 原始材料、公开信息搜索验证、工商数据核查")
            continue
        cleaned.append(_sanitize_text(line))
    return "\n".join(cleaned)


# ── 主入口 ────────────────────────────────────────────

DIMENSION_TITLES = {
    "team": "1. 团队与合规",
    "tech": "2. 技术与产品",
    "industry": "3. 行业与供应链",
    "competition": "4. 竞争与结论",
}

DIMENSION_LABELS = {
    "team": "团队与合规",
    "tech": "技术与产品",
    "industry": "行业与供应链",
    "competition": "竞争与结论",
}


def build_bp_dd_report(
    task_id: str,
    entity: str,
    dimension_outputs: dict[str, str],
    output_path: str,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        _build_bp_dd_markdown(task_id, entity, dimension_outputs, output_path)
        return Path(output_path)

    doc = Document()

    # ── 样式 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    _set_eastasia_font_on_style(style, "宋体")
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.2

    # Heading 样式定制（深蓝色系）
    for level, size_pt, sp_before in [(1, 14, 24), (2, 13, 10), (3, 12, 8)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Microsoft YaHei"
        h_style.font.size = Pt(size_pt)
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        _set_eastasia_font_on_style(h_style, "宋体")
        h_style.paragraph_format.space_before = Pt(sp_before)
        h_style.paragraph_format.space_after = Pt(4)

    # 列表样式
    for list_style_name in ["List Bullet", "List Number"]:
        try:
            ls = doc.styles[list_style_name]
            ls.font.name = "Microsoft YaHei"
            ls.font.size = Pt(10.5)
            _set_eastasia_font_on_style(ls, "宋体")
            ls.paragraph_format.space_after = Pt(3)
        except KeyError:
            pass

    # ── 封面 ──
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("商业计划尽调报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    _set_eastasia_font_on_run(run, "宋体")

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(entity)
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)
    _set_eastasia_font_on_run(run, "宋体")

    doc.add_paragraph("")
    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"生成日期: {time.strftime('%Y年%m月%d日')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("机密 — 仅供内部参考")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run.font.italic = True

    doc.add_page_break()

    # ── 判断模式：统稿 vs 维度拼接 ──
    is_synthesis = "synthesis" in dimension_outputs

    if is_synthesis:
        # ── 统稿模式：整篇报告已经是投研逻辑结构 ──
        content = dimension_outputs["synthesis"]
        content = _clean_dimension_output(content)
        content, all_sources = _strip_source_section(content)

        # 从统稿中提取目录（# 标题行）
        toc_items = []
        for line in content.split("\n"):
            m = re.match(r"^#\s+(.+)$", line.strip())
            if m:
                toc_items.append(m.group(1))

        # 目录
        doc.add_heading("目录", level=1)
        for item in toc_items:
            p = doc.add_paragraph(item)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
        doc.add_page_break()

        # 渲染统稿正文
        _render_markdown_to_doc(doc, content)

    else:
        # ── 维度拼接 fallback 模式 ──
        doc.add_heading("目录", level=1)
        toc_items = [
            "1. 团队与合规",
            "2. 技术与产品",
            "3. 行业与供应链",
            "4. 竞争与结论",
            "5. 综合评估与建议",
            "6. 来源与参考",
            "7. 免责声明",
        ]
        for item in toc_items:
            p = doc.add_paragraph(item)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
        doc.add_page_break()

        all_sources: list[dict] = []

        for slug, dim_title in DIMENSION_TITLES.items():
            content = dimension_outputs.get(slug, "")
            if not content:
                content = "该维度分析未完成。"

            content = _clean_dimension_output(content)
            content, dim_sources = _strip_source_section(content)

            for src in dim_sources:
                src["dimension"] = DIMENSION_LABELS.get(slug, slug)
            all_sources.extend(dim_sources)

            doc.add_heading(dim_title, level=1)

            # 跳过子代理输出的顶部 # 标题
            lines = content.split("\n")
            first_h1_idx = next(
                (idx for idx, l in enumerate(lines) if l.strip().startswith("# ")),
                -1
            )
            if first_h1_idx >= 0:
                lines = lines[first_h1_idx + 1:]
                content = "\n".join(lines)

            _render_markdown_to_doc(doc, content)
            doc.add_page_break()

        # 综合评估（仅维度模式）
        doc.add_heading("5. 综合评估与建议", level=1)
        completed = len(dimension_outputs)
        total = len(DIMENSION_TITLES)
        doc.add_paragraph(f"本次尽调完成 {completed}/{total} 个维度分析。")
        if completed < total:
            missing = [v for k, v in DIMENSION_TITLES.items() if k not in dimension_outputs]
            doc.add_paragraph(f"未完成维度: {', '.join(missing)}")
        doc.add_paragraph("")
        doc.add_page_break()

    # ── 来源与参考（两种模式共用）──
    doc.add_heading("来源与参考", level=1)
    doc.add_paragraph("以下为本报告引用的全部外部来源，统一编号。正文中 [N] 上标标记对应此处编号。")
    doc.add_paragraph("")

    if all_sources:
        # 区分脚注来源和表格来源
        footnote_sources = [s for s in all_sources if s.get("is_footnote")]
        table_sources = [s for s in all_sources if not s.get("is_footnote")]
        
        # 优先渲染脚注格式的来源（按编号排序）
        if footnote_sources:
            # 按编号排序
            try:
                footnote_sources.sort(key=lambda s: int(s.get("id", "0")))
            except (ValueError, TypeError):
                pass
            
            for src in footnote_sources:
                fn_id = src.get("id", "?")
                name = src.get("name", "")
                url = src.get("url", "")
                usage = src.get("usage", "")
                
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                
                # 编号（上标样式）
                run_id = p.add_run(f"[{fn_id}] ")
                run_id.font.bold = True
                run_id.font.size = Pt(10)
                run_id.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
                _set_eastasia_font_on_run(run_id, "宋体")
                
                # 来源名称
                if name:
                    run_name = p.add_run(f"{name}")
                    run_name.font.size = Pt(10)
                    _set_eastasia_font_on_run(run_name, "宋体")
                
                # URL
                if url:
                    if name:
                        run_sep = p.add_run(" — ")
                        run_sep.font.size = Pt(10)
                        run_sep.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                        _set_eastasia_font_on_run(run_sep, "宋体")
                    run_url = p.add_run(url)
                    run_url.font.size = Pt(9)
                    run_url.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
                    _set_eastasia_font_on_run(run_url, "宋体")
                
                # 用途
                if usage:
                    run_usage = p.add_run(f"（{usage}）")
                    run_usage.font.size = Pt(9)
                    run_usage.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    _set_eastasia_font_on_run(run_usage, "宋体")
        
        # 然后渲染表格格式的来源（非脚注）
        if table_sources:
            source_rows = [["编号", "维度", "来源", "URL", "用途"]]
            start_idx = len(footnote_sources) + 1
            for idx, src in enumerate(table_sources, start_idx):
                source_rows.append([
                    f"[{idx}]",
                    src.get("dimension", ""),
                    src.get("name", ""),
                    src.get("url", ""),
                    src.get("usage", ""),
                ])
            _add_table_to_doc(doc, source_rows)
    else:
        # fallback: 从正文提取 URL
        all_urls: list[str] = []
        for slug in DIMENSION_TITLES:
            content = dimension_outputs.get(slug, "")
            urls = re.findall(r"https?://[^\s\)\]\"']+", content)
            all_urls.extend(urls)

        seen: set[str] = set()
        deduped: list[str] = []
        for url in all_urls:
            url = url.rstrip(".,;:)")
            if url not in seen:
                seen.add(url)
                deduped.append(url)

        if deduped:
            source_rows = [["编号", "URL"]]
            for idx, url in enumerate(deduped[:40], 1):
                source_rows.append([f"[{idx}]", url])
            _add_table_to_doc(doc, source_rows)
        else:
            doc.add_paragraph("详见各维度报告中的来源标注。")

    doc.add_page_break()

    # ── 免责声明 ──
    doc.add_heading("7. 免责声明", level=1)
    disclaimers = [
        "本报告由 AI 辅助生成，基于 BP 材料及公开网络搜索结果整理，仅供内部参考。",
        "公开搜索结果仅作线索参考，不等同于工商数据库核验结论。",
        "如进入投资流程，应继续使用工商数据库、法院公告系统、专利数据库全量检索及律师尽调做最终确认。",
        "本报告不构成任何投资建议。投资决策需基于独立判断。",
    ]
    for d in disclaimers:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(d)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _set_eastasia_font_on_run(run, "宋体")

    # ── 保存 ──
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return output


def _build_bp_dd_markdown(
    task_id: str,
    entity: str,
    dimension_outputs: dict[str, str],
    output_path: str,
) -> Path:
    lines = [
        f"# 商业计划尽调报告: {entity}\n",
        f"生成日期: {time.strftime('%Y-%m-%d')}\n",
        "---\n",
    ]
    all_sources = []
    for slug, title in DIMENSION_TITLES.items():
        content = dimension_outputs.get(slug, "**该维度分析未完成。**")
        content = _clean_dimension_output(content)
        content, dim_sources = _strip_source_section(content)
        all_sources.extend(dim_sources)
        lines.append(f"## {title}\n\n{content}\n\n---\n")

    lines.append("## 5. 综合评估与建议\n")
    lines.append(f"本次尽调完成 {len(dimension_outputs)}/{len(DIMENSION_TITLES)} 个维度分析。\n")

    lines.append("\n## 6. 来源与参考\n")
    for idx, src in enumerate(all_sources, 1):
        lines.append(f"[{idx}] {src.get('name', '')} — {src.get('url', '')} ({src.get('usage', '')})")
    lines.append("\n\n*免责声明：本报告由 AI 辅助生成，仅供参考，不构成投资建议。*\n")

    output = Path(output_path)
    if output.suffix == ".docx":
        output = output.with_suffix(".md")
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n".join(lines), encoding="utf-8")
    return output


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Build BP DD Report")
    ap.add_argument("--task-id", required=True)
    ap.add_argument("--entity", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--dimensions-dir", required=True)
    args = ap.parse_args()

    dim_dir = Path(args.dimensions_dir)
    dimension_outputs = {}

    # 优先使用统稿
    synthesis_path = dim_dir / "bp_synthesis.md"
    if synthesis_path.exists():
        dimension_outputs["synthesis"] = synthesis_path.read_text(encoding="utf-8")
    else:
        for slug in ("team", "tech", "industry", "competition"):
            for prefix in ("bp_phase2_", "bp_phase4_"):
                path = dim_dir / f"{prefix}{slug}.md"
                if path.exists():
                    dimension_outputs[slug] = path.read_text(encoding="utf-8")
                    break

    result = build_bp_dd_report(args.task_id, args.entity, dimension_outputs, args.output)
    print(f"Report generated: {result}")
