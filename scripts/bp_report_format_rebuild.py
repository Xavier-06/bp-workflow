#!/usr/bin/env python3
"""
Rebuild BP DD report DOCX from existing DOCX to clean sell-side format.
Reads styles + text from existing DOCX, then builds a clean document with:
- Title page
- Proper heading hierarchy (Word Heading 1-4)
- Real Word tables (no pipe tables)
- Native lists (no checkbox artifacts or raw markdown)
- Source references cleaned up
"""

import sys
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


INPUT_DOCX  = Path("reports/利玛软件有限公司_尽调报告_2026-04-01.docx")
OUTPUT_DOCX = Path("reports/利玛软件有限公司_尽调报告_2026-04-01_rebuilt.docx")

TITLE    = "利玛软件有限公司"
SUBTITLE = "商业尽调报告"
META     = "报告日期：2026年04月01日\n机密文件 — 仅供内部使用"


def _set_font(run, size=10.5, bold=False, italic=False, color=None):
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color
    try:
        rpr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        rFonts.set(qn("w:ascii"),   "Calibri")
        rFonts.set(qn("w:hAnsi"),   "Calibri")
        if rpr.find(qn("w:rFonts")) is not None:
            rpr.remove(rpr.find(qn("w:rFonts")))
        rpr.append(rFonts)
    except Exception:
        pass


def _para_fmt(p, size=10.5, sp_after=6, sp_before=0, align=None, ls=1.15):
    pf = p.paragraph_format
    pf.space_after  = Pt(sp_after)
    pf.space_before = Pt(sp_before)
    pf.line_spacing = ls
    if align is not None:
        p.alignment = align


def _clean_md(text):
    text = re.sub(r"\*\*(.+?)\*\*",  lambda m: m.group(1), text)
    text = re.sub(r"\*(.+?)\*",      lambda m: m.group(1), text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", lambda m: "[" + m.group(1) + "]", text)
    text = re.sub(r"`(.+?)`",        lambda m: m.group(1), text)
    text = re.sub(r"#{1,6}\s+", "", text)
    text = re.sub(r"\|\s+", "", text)
    return text.strip()


def _is_table_sep(text):
    return bool(re.match(r"^[\|\s\-:]+$", text.strip()))


def _is_pipe_table_line(text):
    return "|" in text and text.count("|") >= 3


def extract_blocks(doc):
    blocks = []
    i = 0
    paras = doc.paragraphs

    while i < len(paras):
        p = paras[i]
        raw = p.text.strip()

        if p.style and p.style.name.startswith("Heading"):
            level = 1
            if "2" in p.style.name:   level = 2
            elif "3" in p.style.name: level = 3
            elif "4" in p.style.name: level = 4
            blocks.append(("heading", level, raw))
            i += 1
            continue

        if p.style and p.style.name.startswith("List"):
            ltype = "bullet" if "Bullet" in p.style.name else "number"
            m = re.match(r"^[☐☑□✓✔]\s*(.*)", raw)
            if m:
                blocks.append(("list", m.group(1).strip(), "bullet"))
            else:
                blocks.append(("list", raw, ltype))
            i += 1
            continue

        if _is_pipe_table_line(raw):
            rows = []
            while i < len(paras) and _is_pipe_table_line(paras[i].text.strip()):
                rtext = paras[i].text.strip().strip("|")
                cells = [c.strip() for c in rtext.split("|")]
                rows.append(cells)
                i += 1
            if rows:
                rows = [r for r in rows if not _is_table_sep("|".join(r))]
                blocks.append(("table", rows))
            continue

        cb = re.match(r"^[☐☑□✓✔]\s+(.*)", raw)
        if cb:
            blocks.append(("list", cb.group(1).strip(), "bullet"))
            i += 1
            continue
        bl = re.match(r"^[•·\-–—]\s+(.*)", raw)
        if bl:
            blocks.append(("list", bl.group(1).strip(), "bullet"))
            i += 1
            continue
        nl = re.match(r"^(\d+)\.\s+(.*)", raw)
        if nl:
            blocks.append(("list", nl.group(2).strip(), "number"))
            i += 1
            continue

        if not raw:
            blocks.append(("empty",))
            i += 1
            continue

        blocks.append(("p", raw))
        i += 1
    return blocks


def build_docx(blocks):
    doc = Document()

    for section in doc.sections:
        section.page_width   = Cm(21.0)
        section.page_height  = Cm(29.7)
        section.top_margin   = Cm(2.54)
        section.bottom_margin= Cm(2.54)
        section.left_margin  = Cm(3.18)
        section.right_margin = Cm(3.18)

    # Title page
    p = doc.add_paragraph()
    _para_fmt(p, size=28, sp_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(TITLE)
    _set_font(run, size=28, bold=True)

    p = doc.add_paragraph()
    _para_fmt(p, size=18, sp_after=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(SUBTITLE)
    _set_font(run, size=18, bold=True)

    for ml in META.split("\n"):
        p = doc.add_paragraph()
        _para_fmt(p, size=11, sp_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(ml)
        _set_font(run, size=11)

    doc.add_page_break()

    # TOC
    p = doc.add_paragraph()
    _para_fmt(p, size=16, sp_after=16)
    run = p.add_run("目  录")
    _set_font(run, size=16, bold=True)

    toc_entries = []
    for b in blocks:
        if b[0] == "heading" and b[1] == 1:
            toc_entries.append(b[2])
    for idx, te in enumerate(toc_entries, 1):
        p = doc.add_paragraph()
        _para_fmt(p, size=11, sp_after=3, sp_before=2)
        run = p.add_run(f"{idx}.  {te} ")
        _set_font(run, size=11, bold=True)

    doc.add_page_break()

    heading_sizes = {1: 16, 2: 14, 3: 12, 4: 11}

    def add_para(text, size=10.5, bold=False, italic=False, sp_after=6, sp_before=0, color=None, align=None):
        p = doc.add_paragraph()
        _para_fmt(p, size=size, sp_after=sp_after, sp_before=sp_before, align=align)
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic, color=color)
        return p

    pending_list = []
    pending_type = None

    def flush_list():
        nonlocal pending_list, pending_type
        if not pending_list:
            return
        style_name = "List Bullet" if pending_type == "bullet" else "List Number"
        for txt in pending_list:
            p = doc.add_paragraph(style=style_name)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.line_spacing = Pt(14)
            run = p.add_run(txt)
            _set_font(run, size=10.5)
        pending_list = []
        pending_type = None

    for b in blocks:
        kind = b[0]

        if kind == "heading" and b[2] in (TITLE, SUBTITLE, "商业尽调报告", "目录"):
            continue
        if kind == "p" and any(k in b[1] for k in ["报告日期：2026", "机密文件", "-- 仅供内部使用"]):
            continue
        if kind == "p" and (b[1].startswith("第一部分：") or b[1].startswith("第二部分：") or b[1].startswith("第三部分：") or b[1].startswith("第四部分：") or b[1].startswith("第五部分：") or b[1].startswith("第六部分：")):
            continue

        if kind == "heading":
            flush_list()
            level = min(b[1], 4)
            hsz = heading_sizes[level]
            p = doc.add_paragraph()
            _para_fmt(p, size=hsz, sp_before=18 if level<=2 else 12, sp_after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(b[2])
            _set_font(run, size=hsz, bold=(level<=2))
            if level == 2:
                run.underline = True

        elif kind == "list":
            ltype = b[2]
            if pending_type and pending_type != ltype:
                flush_list()
            pending_list.append(b[1])
            pending_type = ltype

        elif kind == "table":
            flush_list()
            rows = b[1]
            if len(rows) < 2:
                continue
            ncols = max(len(r) for r in rows)
            if ncols < 2:
                continue

            clean = []
            for r in rows:
                cr = []
                for c in r:
                    c = c.strip()
                    c = _clean_md(c)
                    cr.append(c)
                while len(cr) < ncols:
                    cr.append("")
                clean.append(cr)

            table = doc.add_table(rows=len(clean), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for ri, row in enumerate(clean):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    p = cell.paragraphs[0]
                    p.clear()
                    run = p.add_run(cell_text)
                    is_hdr = (ri == 0)
                    tsz = 9 if ncols > 5 else 10
                    _set_font(run, size=tsz, bold=is_hdr)
                    p.paragraph_format.space_after  = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    tc   = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW  = OxmlElement("w:tcW")
                    tcW.set(qn("w:w"),   str(int(9600 // ncols)))
                    tcW.set(qn("w:type"), "dxa")
                    tcPr.append(tcW)

        elif kind == "p":
            flush_list()
            text = _clean_md(b[1])

            if not text or text in ("---", "```", "`"):
                continue
            if "[内部路径已清除]" in text:
                continue
            if text.startswith("[bp_") or text.startswith("file:") or text.startswith("[分析") or text.startswith("[任务"):
                continue
            if any(k in text for k in ["任务 ID:", "生成时间:", "分析完成时间:", "报告行数:", "证据引用:", "完成时间:", "分析师:", "分析师备注:", "报告完成时间:", "报告版本:", "数据来源：", "来源:"]):
                continue

            add_para(text, size=10.5, sp_after=6)

    flush_list()

    # Disclaimer
    doc.add_page_break()
    add_para("免责声明", size=13, bold=True, sp_after=12)
    add_para("本报告仅供内部研究参考使用，不构成任何投资建议。报告中的信息来源于公开渠道，我们不对信息的完整性和准确性做出保证。投资决策需基于进一步的尽职调查和专业判断。本报告内容受保密义务约束，未经授权不得向第三方披露。", size=9, italic=True, color=RGBColor(0x77, 0x77, 0x77), sp_after=6)

    doc.save(str(OUTPUT_DOCX))
    return len(doc.paragraphs)


if __name__ == "__main__":
    src = Document(str(INPUT_DOCX))
    blocks = extract_blocks(src)

    kinds = {}
    for b in blocks:
        k = b[0]
        kinds[k] = kinds.get(k, 0) + 1
    print(f"Extracted {len(blocks)} blocks: {dict(kinds)}")

    if not INPUT_DOCX.exists():
        print(f"ERROR: {INPUT_DOCX} not found")
        sys.exit(1)

    np = build_docx(blocks)
    print(f"Built {np} paragraphs → {OUTPUT_DOCX}")
    sz = OUTPUT_DOCX.stat().st_size
    print(f"File size: {sz:,} bytes ({sz/1024:.1f} KB)")
