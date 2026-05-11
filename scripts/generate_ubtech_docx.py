#!/usr/bin/env python3
"""Generate UBTECH research report DOCX from step8_master.md"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_to_docx(md_path, output_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Parse sections
    lines = content.split('\n')
    current_para = []
    
    for line in lines:
        stripped = line.strip()
        
        # Title (H1)
        if stripped.startswith('# '):
            if current_para:
                doc.add_paragraph('\n'.join(current_para))
                current_para = []
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # H2
        elif stripped.startswith('## '):
            if current_para:
                doc.add_paragraph('\n'.join(current_para))
                current_para = []
            doc.add_heading(stripped[3:], level=2)
        
        # H3
        elif stripped.startswith('### '):
            if current_para:
                doc.add_paragraph('\n'.join(current_para))
                current_para = []
            doc.add_heading(stripped[4:], level=3)
        
        # Table (pipe format)
        elif stripped.startswith('|') and '---' not in stripped:
            # Skip separator lines
            if '---' in line:
                continue
            # Parse table row
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells and any(c for c in cells):
                # Check if we need to create a new table
                if not hasattr(parse_markdown_to_docx, 'current_table'):
                    parse_markdown_to_docx.current_table = None
                
                if parse_markdown_to_docx.current_table is None:
                    # Create new table
                    parse_markdown_to_docx.current_table = doc.add_table(rows=1, cols=len(cells))
                    parse_markdown_to_docx.current_table.style = 'Table Grid'
                    for i, cell_text in enumerate(cells):
                        parse_markdown_to_docx.current_table.rows[0].cells[i].text = cell_text
                else:
                    # Add row to existing table
                    row = parse_markdown_to_docx.current_table.add_row()
                    for i, cell_text in enumerate(cells):
                        if i < len(row.cells):
                            row.cells[i].text = cell_text
        
        # Empty line - end table
        elif stripped == '':
            if current_para:
                doc.add_paragraph('\n'.join(current_para))
                current_para = []
            parse_markdown_to_docx.current_table = None
        
        # Regular text
        elif stripped:
            # Clean markdown artifacts
            cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)  # Remove bold
            cleaned = re.sub(r'\[(\d+)\]', '', cleaned)  # Remove citations
            current_para.append(cleaned)
    
    if current_para:
        doc.add_paragraph('\n'.join(current_para))
    
    doc.save(output_path)
    print(f"Generated: {output_path}")

if __name__ == '__main__':
    parse_markdown_to_docx(
        str(Path(__file__).resolve().parent.parent / 'data/tasks/') + '/TASK-20260330-002-step8_master.md',
        str(Path(__file__).resolve().parent.parent / 'reports/') + '/优必选_09880.HK_深度研报_20260330_v2.docx'
    )
