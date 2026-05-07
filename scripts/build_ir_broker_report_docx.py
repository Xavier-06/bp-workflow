#!/usr/bin/env python3
"""
build_ir_broker_report_docx.py — 将 step8 统稿转为券商风格 Word 研报

v5 (2026-05-07):
  - 字体修复：所有 run 显式设置 font.name="Microsoft YaHei"，
    解决 macOS 上因缺 font.name 导致西文字体回退渲染不一致（歪七扭八）的问题
  - _add_inline_formatted_text() 增加 font_size 参数，统一控制字号
  - 表格、列表、引用块、脚注、来源区全部补上显式 font.name
"""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parent.parent
TASKS = ROOT / 'data' / 'tasks'
REPORTS = ROOT / 'reports'


# ─── 东亚字体设置（与 BP build_bp_dd_report_docx.py 对齐）────────

def _set_eastasia_font_on_style(style, font_name: str):
    """在样式级别设置 eastAsia 字体（解决 MS 明朝问题）"""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _set_eastasia_font_on_run(run, font_name: str):
    """在 run 级别设置 eastAsia 字体"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


# ─── 内部信息清洗模式 ────────────────────────────

INTERNAL_PATTERNS = [
    (r'/Users/\w+/[^\s\n]+', ''),
    (r'/home/\w+/[^\s\n]+', ''),
    (r'~/.openclaw/[^\s\n]+', ''),
    (r'~/.workbuddy/[^\s\n]+', ''),
    (r'data/tasks/[^\s\n]+', ''),
    (r'TASK-\d{8}-\d{3}[-\w]*', ''),
    (r'python3?\s+scripts/[^\n]+', ''),
    (r'(?:bin/yf|/opt/homebrew/bin/ddgs)[^\n]*', ''),
    (r'source\s+[^\n]*python_ssl_env[^\n]*', ''),
    (r'sessions_spawn[^\n]*', ''),
    (r'sessions_send[^\n]*', ''),
    (r'spawn[-_]receipt[^\n]*', ''),
    (r'thinking\s*=\s*high', ''),
    (r'(?:pre-search|presearch)\s*(?:results?)?', ''),
    (r'step\d+_\w+\.md', ''),
    (r'输出文件[：:][^\n]+', ''),
    (r'^[-*]\s*(?:Task|Entity|Accepted evidence|Rounds|Generated)[：:][^\n]*$', ''),
]

LINE_DELETE_PATTERNS = [
    r'^\s*(?:输出文件|Output file|Task ID|任务 ID)\s*[：:]',
    r'^\s*```\s*(?:bash|python|shell)',
    r'^\s*```\s*$',
    r'^\s*(?:子代理|subagent|sub-agent)',
]

_INTERNAL_SOURCE_PATTERNS = [
    r"step\d+_\w+\.md",
    r"bp_ocr_text\.txt",
    r"bp_step0_profile\.(json|md)",
    r"company_verify_report\.(json|md)",
    r"bp_presearch_step_\w+\.md",
    r"ir_presearch",
    r"ir_extract_content",
    r"data/tasks/",
]


def _is_internal_source_row(row_text: str) -> bool:
    for pat in _INTERNAL_SOURCE_PATTERNS:
        if re.search(pat, row_text, re.IGNORECASE):
            return True
    return False


def sanitize_text(text: str) -> str:
    lines = text.split('\n')
    cleaned_lines = []
    in_code_block = False
    code_block_is_internal = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_is_internal = bool(re.search(
                    r'```\s*(?:bash|python|shell)', stripped, re.IGNORECASE
                ))
                continue
            else:
                in_code_block = False
                if code_block_is_internal:
                    code_block_is_internal = False
                    continue
                continue

        if in_code_block and code_block_is_internal:
            continue

        should_delete = False
        for pat in LINE_DELETE_PATTERNS:
            if re.match(pat, stripped, re.IGNORECASE):
                should_delete = True
                break
        if should_delete:
            continue
        cleaned_lines.append(line)

    text = '\n'.join(cleaned_lines)
    for pat, repl in INTERNAL_PATTERNS:
        text = re.sub(pat, repl, text, flags=re.MULTILINE | re.IGNORECASE)
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    return text


def read_text(path: Path) -> str:
    if not path.exists():
        return ''
    return path.read_text(encoding='utf-8', errors='ignore')


# ─── 来源剥离 ────────────────────────────────────

def _strip_source_section(markdown: str) -> tuple[str, list[dict]]:
    """从 Markdown 中剥离来源/参考章节，返回 (cleaned_text, sources)。
    
    同时提取 [^N] 格式的脚注定义。
    """
    lines = markdown.split("\n")
    cleaned_lines = []
    sources: list[dict] = []
    in_source_section = False
    in_source_table = False

    # 第一遍：提取 [^N] 脚注定义
    footnote_defs: dict[str, dict] = {}
    remaining_lines = []

    for line in lines:
        stripped = line.strip()
        fn_match = re.match(r"^\[\^(\d+)\]:\s*(.+)$", stripped)
        if fn_match:
            fn_id = fn_match.group(1)
            fn_content = fn_match.group(2).strip()
            url_match = re.search(r"(https?://[^\s\)\]\"']+)", fn_content)
            url = url_match.group(1) if url_match else ""
            name = re.sub(r"\s*https?://[^\s\)\]\"']+\s*", "", fn_content).strip().rstrip("—–- ")
            footnote_defs[fn_id] = {
                "id": fn_id,
                "name": name or fn_content[:80],
                "url": url,
                "usage": "",
                "is_footnote": True,
            }
            continue
        remaining_lines.append(line)

    # 将脚注定义添加到 sources
    for fn_id in sorted(footnote_defs.keys(), key=int):
        sources.append(footnote_defs[fn_id])

    # 第二遍：处理来源列表章节
    for line in remaining_lines:
        stripped = line.strip()

        if re.match(r"^#{1,3}\s+\d*\.?\s*(?:来源|参考|引用|References|Sources)", stripped, re.IGNORECASE):
            in_source_section = True
            in_source_table = False
            continue

        if in_source_section:
            if stripped.startswith("|"):
                if re.match(r"^\|[\s\-:]+\|", stripped):
                    in_source_table = True
                    continue
                cells = [c.strip() for c in stripped.split("|")]
                if cells and cells[0] == "":
                    cells = cells[1:]
                if cells and cells[-1] == "":
                    cells = cells[:-1]

                if not in_source_table:
                    in_source_table = True
                    continue

                if len(cells) >= 2:
                    row_text = " ".join(cells)
                    if _is_internal_source_row(row_text):
                        continue
                    url = ""
                    for cell in cells:
                        if cell.startswith("http"):
                            url = cell
                            break
                    src = {
                        "id": cells[0] if len(cells) > 0 else "",
                        "name": cells[1] if len(cells) > 1 else "",
                        "url": url or (cells[2] if len(cells) > 2 else ""),
                        "usage": cells[-1] if len(cells) > 3 else "",
                    }
                    if src["url"].startswith("http"):
                        sources.append(src)
                continue

            # [^N] 格式的来源列表
            fn_list_match = re.match(r"^\[\^(\d+)\]\s+(.+)$", stripped)
            if fn_list_match:
                fn_id = fn_list_match.group(1)
                fn_content = fn_list_match.group(2).strip()
                url_match = re.search(r"(https?://[^\s\)\]\"']+)", fn_content)
                url = url_match.group(1) if url_match else ""
                name = re.sub(r"\s*https?://[^\s\)\]\"']+\s*", "", fn_content).strip().rstrip("—–- ")
                if fn_id not in footnote_defs:
                    sources.append({
                        "id": fn_id,
                        "name": name or fn_content[:80],
                        "url": url,
                        "usage": "",
                        "is_footnote": True,
                    })
                continue

            # 非表格的来源行（如 [1] xxx — url）
            ref_match = re.match(r"^\[(\d+)\]\s*(.+?)(?:\s*[—–-]\s*(https?://\S+))?(?:\s*\((.+?)\))?\s*$", stripped)
            if ref_match:
                url = ref_match.group(3) or ""
                if url and not _is_internal_source_row(stripped):
                    sources.append({
                        "id": f"[{ref_match.group(1)}]",
                        "name": ref_match.group(2),
                        "url": url,
                        "usage": ref_match.group(4) or "",
                    })
                continue

            if not stripped or re.match(r"^-{3,}$", stripped):
                continue

            if stripped.startswith("#"):
                in_source_section = False
                cleaned_lines.append(line)
                continue
            continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines), sources


# ─── 表格美化（与 BP v3 统一）────────────────────

def parse_markdown_table(lines: list[str], start_idx: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start_idx
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped.startswith('|'):
            break
        if re.match(r'^\|[\s\-:|]+\|$', stripped):
            i += 1
            continue
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def _add_inline_formatted_text(paragraph, text: str, font_size: Pt = Pt(11)):
    """解析行内 Markdown 格式并添加到段落。支持 [^N] 脚注标记渲染为上标。
    
    所有 run 显式设置 font.name="Microsoft YaHei" + font_size，
    避免 macOS 上因缺字体声明导致渲染不一致。
    """
    pattern = r"(\*\*(.+?)\*\*|\[([^\]]+)\]\(([^)]+)\)|\[\^(\d+)\]|`([^`]+)`)"
    last_end = 0
    for m in re.finditer(pattern, text):
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            run.font.name = "Microsoft YaHei"
            run.font.size = font_size
            _set_eastasia_font_on_run(run, "宋体")
        if m.group(2) is not None:
            # **bold**
            run = paragraph.add_run(m.group(2))
            run.font.bold = True
            run.font.name = "Microsoft YaHei"
            run.font.size = font_size
            _set_eastasia_font_on_run(run, "宋体")
        elif m.group(3) is not None and m.group(5) is None:
            # [text](url) hyperlink
            run = paragraph.add_run(m.group(3))
            run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
            run.font.underline = True
            run.font.name = "Microsoft YaHei"
            run.font.size = font_size
            _set_eastasia_font_on_run(run, "宋体")
        elif m.group(5) is not None:
            # [^N] footnote reference → render as superscript
            fn_num = m.group(5)
            run = paragraph.add_run(f"[{fn_num}]")
            run.font.superscript = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
            run.font.name = "Microsoft YaHei"
            _set_eastasia_font_on_run(run, "宋体")
        elif m.group(6) is not None:
            # `code`
            run = paragraph.add_run(m.group(6))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        last_end = m.end()
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = "Microsoft YaHei"
        run.font.size = font_size
        _set_eastasia_font_on_run(run, "宋体")


def add_table_to_doc(doc: Document, rows: list[list[str]]):
    """美化表格：表头底纹、斑马纹、自适应列宽、紧凑内边距。"""
    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 自动布局
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    tblLayout = tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "autofit"})
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    tblPr.append(tblLayout)

    tblW = tblPr.makeelement(qn("w:tblW"), {qn("w:w"): "5000", qn("w:type"): "pct"})
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblPr.append(tblW)

    # 单元格内边距
    tblCellMar = tblPr.makeelement(qn("w:tblCellMar"), {})
    for side, val in [("top", "40"), ("bottom", "40"), ("left", "80"), ("right", "80")]:
        el = tblCellMar.makeelement(qn(f"w:{side}"), {qn("w:w"): val, qn("w:type"): "dxa"})
        tblCellMar.append(el)
    for existing in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(existing)
    tblPr.append(tblCellMar)

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= n_cols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            _add_inline_formatted_text(p, cell_text.strip(), font_size=Pt(9))

            for run in p.runs:
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(9)
                _set_eastasia_font_on_run(run, "宋体")

            if i == 0:
                for run in p.runs:
                    run.font.bold = True
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="1F4E79" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                if i % 2 == 0:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F2F7FB" w:val="clear"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading)

    p = doc.add_paragraph("")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# ─── Markdown → DOCX 转换（完整版）──────────────

def convert_markdown_to_docx(text: str, doc: Document):
    lines = text.split('\n')
    i = 0
    in_code_block = False

    while i < len(lines):
        stripped = lines[i].strip()

        # 代码块
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 分隔线
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {
                qn("w:val"): "single", qn("w:sz"): "4",
                qn("w:space"): "1", qn("w:color"): "CCCCCC",
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 表格
        if stripped.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            table_rows, end_idx = parse_markdown_table(lines, i)
            if table_rows:
                add_table_to_doc(doc, table_rows)
            i = end_idx
            continue

        # 标题
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        # 引用块
        if stripped.startswith("> "):
            quote_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            _add_inline_formatted_text(p, quote_text, font_size=Pt(10))
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.italic = True
                # font.name and font.size already set by _add_inline_formatted_text
            i += 1
            continue

        # 无序列表
        if stripped.startswith('- ') or stripped.startswith('* '):
            item_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_formatted_text(p, item_text, font_size=Pt(10.5))
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if ol_match:
            item_text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            _add_inline_formatted_text(p, item_text, font_size=Pt(10.5))
            i += 1
            continue

        # 纯加粗行
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.font.bold = True
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(11)
            _set_eastasia_font_on_run(run, "宋体")
            i += 1
            continue

        # 普通段落
        if len(stripped) > 2:
            p = doc.add_paragraph()
            _add_inline_formatted_text(p, stripped)
            i += 1
            continue

        i += 1


# ─── 主入口 ──────────────────────────────────────

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('task_id')
    ap.add_argument('--output')
    args = ap.parse_args()

    tid = args.task_id
    REPORTS.mkdir(parents=True, exist_ok=True)
    out_path = Path(args.output) if args.output else REPORTS / f'{tid}-券商版研报.docx'

    pkg_path = TASKS / f'{tid}.json'
    if pkg_path.exists():
        pkg = json.loads(pkg_path.read_text(encoding='utf-8'))
    else:
        registry_path = ROOT / 'tasks' / 'task_registry' / f'{tid}.json'
        if registry_path.exists():
            pkg = json.loads(registry_path.read_text(encoding='utf-8'))
        else:
            pkg = {'query': '深度研报', 'entity': ''}
            step8_fallback = TASKS / f'{tid}-step8_master.md'
            if step8_fallback.exists():
                first_line = step8_fallback.read_text(encoding='utf-8').split('\n', 1)[0]
                m = re.match(r'^#\s+(.+?)(?:深度|投资|研究|研报)', first_line)
                if m:
                    pkg['entity'] = m.group(1).strip()
    query = pkg.get('query', '深度研报')
    entity = pkg.get('entity', '')

    # Read memo
    memo_path = TASKS / f'{tid}-step8_master.md'
    if not memo_path.exists():
        ws_memo = ROOT / 'jobs' / tid / 'outputs' / 'step8_master.md'
        if ws_memo.exists():
            memo_path = ws_memo
        else:
            import sys
            print(f"ERROR: step8_master missing at {memo_path}")
            print("REFUSING to generate half-baked DOCX")
            sys.exit(1)

    memo = read_text(memo_path)
    if not memo or len(memo) < 1000:
        import sys
        print(f"ERROR: step8_master content too short ({len(memo) if memo else 0} chars)")
        print("REFUSING to generate half-baked DOCX")
        sys.exit(1)

    # 硬检查
    if not memo or len(memo) < 200:
        step_files = list(TASKS.glob(f'{tid}-step*.md'))
        step_count = len([f for f in step_files if f.stat().st_size > 200 and '-brief-' not in f.name and '-followup' not in f.name])
        raise RuntimeError(
            f"Phase 5 拦截：step8_master 统稿缺失（{len(memo) if memo else 0} 字符），"
            f"有效 step 产出仅 {step_count} 个。"
            f"禁止生成半成品 DOCX。请先完成 Phase 4 全部 8 个 step。"
        )

    # 清洗内部信息
    memo = sanitize_text(memo)

    # 剥离来源章节
    memo, sources = _strip_source_section(memo)

    doc = Document()

    # ── 样式（与 BP 对齐：font.name = Microsoft YaHei, eastAsia = 宋体）──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    _set_eastasia_font_on_style(style, "宋体")
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.2

    # Heading 样式定制（深蓝色系，与悦享资本报告对齐）
    for level, size_pt, sp_before in [(1, 14, 24), (2, 13, 10), (3, 12, 8)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Microsoft YaHei"
        h_style.font.size = Pt(size_pt)
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        _set_eastasia_font_on_style(h_style, "宋体")
        h_style.paragraph_format.space_before = Pt(sp_before)
        h_style.paragraph_format.space_after = Pt(4)

    # 列表样式
    for list_style_name in ["List Bullet", "List Number"]:
        try:
            ls = doc.styles[list_style_name]
            ls.font.name = "Microsoft YaHei"
            ls.font.size = Pt(10.5)
            _set_eastasia_font_on_style(ls, "宋体")
            ls.paragraph_format.space_after = Pt(3)
        except KeyError:
            pass

    # ── 封面 ──
    for _ in range(4):
        doc.add_paragraph("")

    company_hint = entity or query.replace('深度研报', '').replace('（', '(').replace('）', ')').strip()

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(f'{company_hint} 深度研究报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    _set_eastasia_font_on_run(run, "宋体")

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = meta.add_run(f'生成日期: {datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _set_eastasia_font_on_run(run, "宋体")

    conf = doc.add_paragraph()
    conf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = conf.add_run('内部研究讨论稿 — 非投资建议')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run.font.italic = True
    _set_eastasia_font_on_run(run, "宋体")

    doc.add_page_break()

    # ── 主体内容 ──
    convert_markdown_to_docx(memo, doc)

    # ── 来源与参考（统一放末尾）──
    doc.add_page_break()
    doc.add_heading('来源与参考', level=1)

    if sources:
        # 区分脚注来源和表格来源
        footnote_sources = [s for s in sources if s.get("is_footnote")]
        table_sources = [s for s in sources if not s.get("is_footnote")]

        if footnote_sources:
            try:
                footnote_sources.sort(key=lambda s: int(s.get("id", "0")))
            except (ValueError, TypeError):
                pass

            doc.add_paragraph("以下为本报告引用的全部外部来源，统一编号。正文中 [N] 上标标记对应此处编号。")
            doc.add_paragraph("")

            for src in footnote_sources:
                fn_id = src.get("id", "?")
                name = src.get("name", "")
                url = src.get("url", "")

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

                run_id = p.add_run(f"[{fn_id}] ")
                run_id.font.bold = True
                run_id.font.name = "Microsoft YaHei"
                run_id.font.size = Pt(10)
                run_id.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
                _set_eastasia_font_on_run(run_id, "宋体")

                if name:
                    run_name = p.add_run(f"{name}")
                    run_name.font.name = "Microsoft YaHei"
                    run_name.font.size = Pt(10)
                    _set_eastasia_font_on_run(run_name, "宋体")

                if url:
                    if name:
                        run_sep = p.add_run(" — ")
                        run_sep.font.name = "Microsoft YaHei"
                        run_sep.font.size = Pt(10)
                        run_sep.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                        _set_eastasia_font_on_run(run_sep, "宋体")
                    run_url = p.add_run(url)
                    run_url.font.name = "Microsoft YaHei"
                    run_url.font.size = Pt(9)
                    run_url.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
                    _set_eastasia_font_on_run(run_url, "宋体")

        if table_sources:
            source_rows = [["编号", "来源", "URL", "用途"]]
            start_idx = len(footnote_sources) + 1
            for idx, src in enumerate(table_sources, start_idx):
                source_rows.append([
                    f"[{idx}]",
                    src.get("name", ""),
                    src.get("url", ""),
                    src.get("usage", ""),
                ])
            add_table_to_doc(doc, source_rows)
    else:
        # fallback: 从正文提取 URL
        urls = re.findall(r"https?://[^\s\)\]\"']+", memo)
        seen: set[str] = set()
        deduped: list[str] = []
        for url in urls:
            url = url.rstrip(".,;:)")
            if url not in seen:
                seen.add(url)
                deduped.append(url)
        if deduped:
            source_rows = [["编号", "URL"]]
            for idx, url in enumerate(deduped[:40], 1):
                source_rows.append([f"[{idx}]", url])
            add_table_to_doc(doc, source_rows)
        else:
            doc.add_paragraph("详见正文中的来源标注。")

    # ── 免责声明 ──
    doc.add_page_break()
    doc.add_heading('免责声明', level=1)
    disclaimers = [
        '本报告由 AI 系统辅助生成，仅供内部研究讨论使用，不构成任何投资建议。',
        '报告中的数据来源于公开信息，可能存在滞后或偏差。',
        '使用者应自行核实关键数据并独立做出投资决策。',
    ]
    for d in disclaimers:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(d)
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _set_eastasia_font_on_run(run, "宋体")

    # ── 保存 ──
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(json.dumps({
        'task_id': tid,
        'output': str(out_path),
        'paragraphs': len(doc.paragraphs),
        'tables': len(doc.tables),
        'sources_extracted': len(sources),
        'sanitized': True,
    }, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
