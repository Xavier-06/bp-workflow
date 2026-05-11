#!/usr/bin/env python3
"""
投研报告生成脚本
支持 Excel 和 Word 格式，自动生成行业分析报告、公司调研报告等
"""
import sys
import os
from datetime import datetime
from pathlib import Path

# ============================================================
#  Excel 报告生成
# ============================================================

def generate_excel_report(title, data, output_path):
    """
    生成 Excel 格式的投研报告
    
    Args:
        title: 报告标题
        data: 字典格式的数据，包含多个 sheet
        output_path: 输出文件路径
    """
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, Color
    from openpyxl.utils import get_column_letter
    
    # 创建工作簿
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # 移除默认 sheet
    
    # 样式定义
    title_font = Font(name='微软雅黑', size=18, bold=True, color='FFFFFF')
    header_font = Font(name='微软雅黑', size=12, bold=True, color='FFFFFF')
    normal_font = Font(name='微软雅黑', size=11)
    
    title_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    header_fill = PatternFill(start_color='2E75B6', end_color='2E75B6', fill_type='solid')
    even_fill = PatternFill(start_color='D6EAF8', end_color='D6EAF8', fill_type='solid')
    
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # 创建封面页
    ws_cover = wb.create_sheet('封面')
    ws_cover.merge_cells('A1:D4')
    cover_cell = ws_cover['A1']
    cover_cell.value = title
    cover_cell.font = title_font
    cover_cell.fill = title_fill
    cover_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws_cover.row_dimensions[1].height = 80
    
    # 生成日期
    ws_cover['A6'] = f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ws_cover['A6'].font = normal_font
    
    # 创建数据 sheets
    for sheet_name, sheet_data in data.items():
        ws = wb.create_sheet(sheet_name)
        
        # 写入表头
        if 'headers' in sheet_data:
            headers = sheet_data['headers']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.border = thin_border
            
            # 写入数据
            rows = sheet_data.get('rows', [])
            for row_idx, row in enumerate(rows, 2):
                for col_idx, value in enumerate(row, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.font = normal_font
                    cell.border = thin_border
                    # 隔行填色
                    if row_idx % 2 == 0:
                        cell.fill = even_fill
            
            # 调整列宽
            for col_idx, header in enumerate(headers, 1):
                col_letter = get_column_letter(col_idx)
                ws.column_dimensions[col_letter].width = max(len(str(header)) + 2, 15)
        
        # 添加汇总统计（如果有数值数据）
        if 'stats' in sheet_data:
            stats_row = len(sheet_data.get('rows', [])) + 3
            for stat_label, stat_value in sheet_data['stats'].items():
                ws.cell(row=stats_row, column=1, value=stat_label).font = Font(bold=True)
                ws.cell(row=stats_row, column=2, value=stat_value)
                stats_row += 1
    
    # 保存文件
    wb.save(output_path)
    print(f"✅ Excel 报告已生成：{output_path}")
    return output_path


# ============================================================
#  Word 报告生成
# ============================================================

def generate_word_report(title, sections, output_path):
    """
    生成 Word 格式的投研报告
    
    Args:
        title: 报告标题
        sections: 列表，每个元素是 {'heading': str, 'content': str, 'table': dict} 的字典
        output_path: 输出文件路径
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    
    doc = Document()
    
    # 标题样式
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)  # 深蓝色
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题 - 生成时间
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 添加各个章节
    for section in sections:
        # 章节标题
        if 'heading' in section:
            heading = doc.add_heading(section['heading'], level=1)
            heading.runs[0].font.color.rgb = RGBColor(46, 117, 182)
        
        # 章节内容
        if 'content' in section:
            content_para = doc.add_paragraph(section['content'])
            content_para.runs[0].font.size = Pt(11)
        
        # 添加表格（如果有）
        if 'table' in section:
            table_data = section['table']
            headers = table_data.get('headers', [])
            rows = table_data.get('rows', [])
            
            if headers and rows:
                table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                table.style = 'Table Grid'
                
                # 表头
                header_row = table.rows[0]
                for i, header in enumerate(headers):
                    cell = header_row.cells[i]
                    cell.text = header
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 数据行
                for row_idx, row_data in enumerate(rows, 1):
                    for col_idx, value in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = str(value)
        
        doc.add_paragraph()  # 章节间空行
    
    # 保存文件
    doc.save(output_path)
    print(f"✅ Word 报告已生成：{output_path}")
    return output_path


# ============================================================
#  命令行接口
# ============================================================

def main():
    if len(sys.argv) < 2:
        print("投研报告生成工具")
        print("=" * 50)
        print("\n用法:")
        print("  python3 generate_report.py <report_type> [options]")
        print("\n支持的报告类型:")
        print("  excel   - 生成 Excel 格式报告")
        print("  word    - 生成 Word 格式报告")
        print("  demo    - 生成示例报告（Excel + Word）")
        print("\n示例:")
        print("  python3 generate_report.py demo")
        sys.exit(1)
    
    report_type = sys.argv[1]
    
    if report_type == "demo":
        # 生成示例报告
        output_dir = Path(__file__).resolve().parent.parent / 'reports'
        output_dir.mkdir(exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Excel 示例
        excel_data = {
            "行业数据": {
                "headers": ["指标", "2024 年", "2025 年", "同比增长"],
                "rows": [
                    ["投资数量", "3,964 起", "5,074 起", "28%"],
                    ["投资规模", "4,872 亿元", "5,748 亿元", "18%"],
                    ["平均单笔", "1.23 亿元", "1.13 亿元", "-8%"],
                ],
                "stats": {
                    "数据来源": "沙利文新投资大会",
                    "更新时间": "2025-06"
                }
            },
            "热门赛道": {
                "headers": ["赛道", "投资事件", "占比", "趋势"],
                "rows": [
                    ["人工智能", "1,218 起", "24%", "↑"],
                    ["半导体", "892 起", "18%", "↑"],
                    ["新能源", "761 起", "15%", "→"],
                    ["生物医药", "654 起", "13%", "↓"],
                ]
            }
        }
        
        excel_path = output_dir / f"PEVC_行业分析报告_{timestamp}.xlsx"
        generate_excel_report("2025 年中国 PE/VC 行业分析报告", excel_data, excel_path)
        
        # Word 示例
        word_sections = [
            {
                "heading": "市场概览",
                "content": "2025 年上半年，中国 PE/VC 市场投资呈现回暖态势，投资数量共 5,074 起，同比增长 28%；投资总规模达 5,748 亿元，同比增长 18%。硬科技领域成为投资焦点，人工智能、半导体等行业获得大量资金青睐。"
            },
            {
                "heading": "投资趋势",
                "content": "从投资赛道分布来看，人工智能以 1,218 起投资事件位居首位，占比 24%；半导体行业紧随其后，投资事件 892 起，占比 18%。新能源和生物医药分别占比 15% 和 13%。",
                "table": {
                    "headers": ["赛道", "投资事件", "占比"],
                    "rows": [
                        ["人工智能", "1,218 起", "24%"],
                        ["半导体", "892 起", "18%"],
                        ["新能源", "761 起", "15%"],
                    ]
                }
            },
            {
                "heading": "政策环境",
                "content": "政策红利及产业需求升级驱动中国 PE/VC 市场发展。2025 年以来，多项支持科技创新和产业升级的政策出台，为投资机构提供了良好的政策环境。"
            }
        ]
        
        word_path = output_dir / f"PEVC_行业分析报告_{timestamp}.docx"
        generate_word_report("2025 年中国 PE/VC 行业分析报告", word_sections, word_path)
        
        print(f"\n📊 报告已生成到：{output_dir}")
        print(f"   - Excel: {excel_path.name}")
        print(f"   - Word:  {word_path.name}")
        
    else:
        print(f"未知报告类型：{report_type}")
        print("使用 'python3 generate_report.py' 查看帮助")
        sys.exit(1)


if __name__ == "__main__":
    main()
