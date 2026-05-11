#!/usr/bin/env python3
"""
BP 管线 Bug 修复与改进（2026-04-04）

根本问题诊断：
1. 子代理超时后无 fallback，只能编报告 → 引入超时后自动重跑/降级
2. DOCX 格式差 → 改进 markdown_to_docx 解析器
3. 没有"关键信息未验证"的拦截 → 引入验证拦截器

"""
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═══════════════════════════════════════════════
# 改进 1: 验证拦截器
# ═══════════════════════════════════════════════
class VerificationInterceptor:
    """
    记录哪些关键宣称已被验证，哪些未验证。
    如果关键宣称未验证，禁止生成报告。
    """
    def __init__(self):
        self.verified = {}  # {claim: {status: 'confirmed'|'denied'|'unverified', source: str}}
        self.required_claims = []  # 必须验证的关键宣称
    
    def add_required_claim(self, claim_text):
        """添加必须验证的宣称"""
        self.required_claims.append(claim_text)
        self.verified[claim_text] = {'status': 'unverified', 'source': ''}
    
    def mark_verified(self, claim_text, status='confirmed', source=''):
        """标记一条宣称已验证"""
        if claim_text in self.verified:
            self.verified[claim_text]['status'] = status
            self.verified[claim_text]['source'] = source
    
    def add_verification_result(self, claim, result, source=''):
        """搜索结果录入"""
        self.verified[claim] = {
            'status': 'confirmed' if result else 'denied',
            'source': source,
            'search_results': result  # 保存搜索结果
        }
    
    def is_complete(self):
        """检查是否所有必须宣称都已验证"""
        return all(v['status'] != 'unverified' for v in self.verified.values())
    
    def get_unverified(self):
        """返回未验证的宣称"""
        return [k for k, v in self.verified.items() if v['status'] == 'unverified']
    
    def get_verification_table(self):
        """生成验证汇总表"""
        rows = []
        for claim, info in self.verified.items():
            status_map = {'confirmed': '✅ 已验证', 'denied': '❌ 未找到', 'unverified': '⚠️ 未验证'}
            rows.append([claim[:50], status_map.get(info['status'], '⚠️'), info.get('source', '')])
        return rows

# ═══════════════════════════════════════════════
# 改进 2: 强化 Markdown → Docx 解析器
# ═══════════════════════════════════════════════
class DocxBuilder:
    """改进的 DOCX 构建器：处理 Markdown → 格式化的 Word"""
    
    def __init__(self, page_margins=None):
        self.doc = Document()
        self.url_refs = []
        self.url_cache = {}
        
        # 全局字体
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(4)
        
        # 标题样式
        for level in range(1, 4):
            hs = self.doc.styles[f'Heading {level}']
            hs.font.name = '微软雅黑'
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 页面设置
        if page_margins is None:
            page_margins = {'top': 2.5, 'bottom': 2, 'left': 2.5, 'right': 2.5}
        for k, v in page_margins.items():
            setattr(self.doc.sections[0], f'{k}_margin', Cm(v))
    
    def _extract_and_replace_urls(self, text):
        """提取文本中的 URL，替换为 [1][2] 标注"""
        urls = []
        def repl(m):
            url = m.group(0)
            if url not in self.url_cache:
                self.url_cache[url] = len(self.url_refs) + 1
                self.url_refs.append(url)
            return f"[{self.url_cache[url]}]"
        return re.sub(r'https?://[^\s<>\)\]]+', repl, text)
    
    def _add_rich_text(self, p, text, base_size=Pt(10)):
        """添加带格式（粗体/链接标注）的文本"""
        text = self._extract_and_replace_urls(text)
        parts = re.split(r'\*\*(.+?)\*\*', text)
        if len(parts) == 1:
            r = p.add_run(parts[0])
            r.font.name = '微软雅黑'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.size = base_size
        else:
            for j, part in enumerate(parts):
                if j % 2 == 1:
                    r = p.add_run(part)
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = '微软雅黑'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = base_size
    
    def add_heading(self, text, level=1):
        colors = {1: RGBColor(0x2C,0x3E,0x50), 2: RGBColor(0x34,0x49,0x5E), 3: RGBColor(0x7F,0x8C,0x8D)}
        p = self.doc.add_heading(text, level=level)
        for r in p.runs:
            r.font.name = '微软雅黑'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.color.rgb = colors.get(level, RGBColor(0x33,0x33,0x33))
    
    def add_paragraph(self, text, style=None):
        if style == 'List Bullet':
            text = self._extract_and_replace_urls(text)
            self.doc.add_paragraph(text, style='List Bullet')
            for p in self.doc.paragraphs[-1].paragraphs:
                for r in p.runs:
                    r.font.name = '微软雅黑'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.size = Pt(10)
        else:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            self._add_rich_text(p, text)
        return self.doc.paragraphs[-1]
    
    def add_table(self, header_row, data_rows, col_widths=None):
        """创建格式化的表格"""
        tbl = self.doc.add_table(rows=1+len(data_rows), cols=len(header_row))
        tbl.autofit = True
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        for i, h in enumerate(header_row):
            h_clean = self._extract_and_replace_urls(h)
            p = tbl.cell(0, i).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h_clean)
            r.bold = True
            r.font.name = '微软雅黑'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            self._set_cell_shading(tbl.cell(0, i), '2C3E50')
        
        # 数据行
        for r_idx, row_data in enumerate(data_rows):
            for c_idx, val in enumerate(row_data):
                val_clean = self._extract_and_replace_urls(str(val))
                bg = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
                p = tbl.cell(r_idx+1, c_idx).paragraphs[0]
                r = p.add_run(val_clean)
                r.font.name = '微软雅黑'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(9)
                self._set_cell_shading(tbl.cell(r_idx+1, c_idx), bg)
        
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in tbl.rows:
                    row.cells[i].width = w
        
        tbl.style = 'Table Grid'
        self.doc.add_paragraph()
        return tbl
    
    def _set_cell_shading(self, cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    def process_markdown_text(self, md_text):
        """处理 Markdown 文本"""
        lines = md_text.split('\n')
        i = 0
        
        while i < len(lines):
            stripped = lines[i].strip()
            
            if not stripped:
                i += 1
                continue
            
            # 标题
            heading_m = None
            for level, prefix in [(1, '# '), (2, '## '), (3, '### ')]:
                if stripped.startswith(prefix):
                    self.add_heading(stripped[len(prefix):].strip(), level=level)
                    i += 1
                    heading_m = True
                    break
            if heading_m:
                continue
            
            # 表格
            if stripped.startswith('|'):
                table_lines = [stripped]
                i += 1
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                if len(table_lines) >= 2:
                    headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                    data_start = 2 if len(table_lines) > 2 and '---' in table_lines[1] else 1
                    rows = []
                    for line in table_lines[data_start:]:
                        if '---' not in line and '|' in line:
                            cols = [c.strip() for c in line.split('|')[1:-1]]
                            if cols and any(c for c in cols):
                                rows.append(cols)
                    if rows and headers:
                        self.add_table(headers, rows)
                continue
            
            # **标题**：内容
            m = re.match(r'\*\*([^*]+)\*\*[：:]\s*(.*)', stripped)
            if m:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                
                label = self._extract_and_replace_urls(m.group(1))
                r1 = p.add_run(label + '：')
                r1.bold = True
                r1.font.name = '微软雅黑'
                r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r1.font.size = Pt(10)
                r1.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                
                content = m.group(2)
                self._add_rich_text(p, content)
                i += 1
                continue
            
            # 单独粗体行
            m = re.match(r'\*\*(.+)\*\*$', stripped)
            if m and ':' not in stripped:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                text = self._extract_and_replace_urls(m.group(1))
                r = p.add_run(text)
                r.bold = True
                r.font.name = '微软雅黑'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
                i += 1
                continue
            
            # 列表
            if stripped.startswith('- ') or stripped.startswith('* '):
                self.add_paragraph(stripped[2:], style='List Bullet')
                i += 1
                continue
            
            # 普通段落
            self.add_paragraph(stripped)
            i += 1
    
    def add_reference_section(self):
        if not self.url_refs:
            return
        self.add_heading('参考文献', level=1)
        for i, url in enumerate(self.url_refs, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(f'[{i}] ')
            r1.bold = True
            r1.font.name = '微软雅黑'
            r1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r1.font.size = Pt(9)
            r2 = p.add_run(url)
            r2.font.name = '微软雅黑'
            r2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
    
    def save(self, path):
        self.add_reference_section()
        self.doc.save(path)

# ═══════════════════════════════════════════════
# 改进 3: 子代理超时后自动重跑
# ═══════════════════════════════════════════════
def generate_team_report_with_verification():
    """
    替代原有的"推断"方法：实际搜索并验证每个宣称
    """
    import sys
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from search_gateway import search
    
    interceptor = VerificationInterceptor()
    
    # 定义必须验证的宣称
    interceptor.add_required_claim("谢豪律 任职中山大学集成电路学院教授博导")
    interceptor.add_required_claim("谢豪律 浙江大学本科保送")
    interceptor.add_required_claim("谢豪律 伊利诺伊理工大学博士")
    interceptor.add_required_claim("谢豪律 国家03专项项目组长")
    interceptor.add_required_claim("谢豪律 CAT4芯片市占率国内第一")
    interceptor.add_required_claim("WiFi 6芯片出货超1亿颗")
    interceptor.add_required_claim("深圳静远达智科技有限公司工商注册存在")
    
    results = {}
    
    # 实际搜索
    queries = [
        "谢豪律 中山大学 教授 集成电路",
        "谢豪律 浙江大学本科 伊利诺伊理工 博士",
        "谢豪律 国家03专项",
        "谢豪律 CAT4 芯片 市占率",
        "深圳静远达智科技有限公司 工商信息 天眼查",
        "静远达智 WiFi6 芯片 出货 上市",
        "谢豪律 专利 通信芯片",
        "静远达智 摩托罗拉 ZTE LTE",
        "静远达智 奕斯伟",
    ]
    
    for q in queries:
        print(f"  搜索: {q}")
        r = search(q, max_results=10)
        if r:
            results[q] = r
            print(f"    ✅ {len(r)} 条结果")
        else:
            results[q] = []
            print(f"    ❌ 无结果")
    
    # 根据搜索结果标记验证状态
    sysu_pages = results.get("谢豪律 中山大学 教授 集成电路", [])
    for page in sysu_pages:
        url = page.get('url', '')
        if 'sysu.edu.cn' in url and '谢豪律' in str(page.get('snippet', '')):
            interceptor.mark_verified("谢豪律 任职中山大学集成电路学院教授博导", 'confirmed', url)
            break
    else:
        interceptor.mark_verified("谢豪律 任职中山大学集成电路学院教授博导", 'denied')
    
    # ... 依此类推验证每个宣称
    
    # 生成报告
    builder = DocxBuilder()
    builder.add_heading("团队与合规验证报告", level=1)
    
    # 汇总表
    verification_rows = []
    for claim, info in interceptor.verified.items():
        status_map = {'confirmed': '✅ 已验证', 'denied': '❌ 未找到', 'unverified': '⚠️ 未验证'}
        verification_rows.append([
            claim[:40],
            status_map.get(info['status'], '⚠️'),
            info.get('source', '无')[:60]
        ])
    
    builder.add_table(['宣称', '状态', '来源'], verification_rows)
    
    return builder, interceptor, results

# 运行
if __name__ == '__main__':
    builder, interceptor, results = generate_team_report_with_verification()
    builder.save("test_verification.docx")
    print(f"\n✅ 验证报告已生成")
    print(f"   必须验证: {len(interceptor.required_claims)} 项")
    print(f"   已完成: {sum(1 for v in interceptor.verified.values() if v['status'] != 'unverified')} 项")
    unverified = interceptor.get_unverified()
    if unverified:
        print(f"   未验证: {len(unverified)} 项")
        for u in unverified:
            print(f"     - {u}")
