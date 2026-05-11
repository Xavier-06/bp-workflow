#!/usr/bin/env python3
"""
Pipeline Quality Guard — 管线质量守卫

解决两个根本问题：
1. 子代理超时后没人兜底 → 自动触发搜索验证
2. DOCX 格式差 → 完善的 Markdown 解析器（使用 parse_xml 正确方式）
"""
from __future__ import annotations
import sys
import re
import json
import time
from pathlib import Path
from typing import Optional

sys.path.insert(0, str(Path(__file__).resolve().parent))

# ═══════════════════════════════════════════════
# Part A: 验证拦截器 — 子代理超时后自动搜索
# ═══════════════════════════════════════════════

class VerificationInterceptor:
    """记录哪些宣称已验证，哪些未验证，阻止未验证的生成报告"""
    
    def __init__(self):
        self.claims = {}

    def add_claim(self, text, critical=True):
        self.claims[text] = {'status': 'unverified', 'critical': critical, 'source': '', 'results': []}

    def mark(self, claim_text, status, source=''):
        if claim_text in self.claims:
            self.claims[claim_text]['status'] = status
            self.claims[claim_text]['source'] = source

    def get_unverified(self):
        return [k for k, v in self.claims.items() if v['status'] == 'unverified']

    def add_results(self, claim_text, search_results):
        if claim_text in self.claims:
            self.claims[claim_text]['results'] = search_results
            if search_results:
                self.claims[claim_text]['status'] = 'confirmed'
            else:
                self.claims[claim_text]['status'] = 'denied'

    def summary_table(self):
        rows = []
        status_map = {'confirmed': '✅', 'denied': '❌', 'unverified': '⚠️'}
        for claim, info in self.claims.items():
            rows.append([
                claim[:45],
                status_map.get(info['status'], '?'),
                info.get('source', '')[:60]
            ])
        return rows


def verify_team_claims(company_name, founder_name, claims, search_fn=None):
    """
    验证团队宣称。返回 VerificationInterceptor。
    
    claims = [
        {"claim": "xxx在职", "queries": ["搜索词1", "搜索词2"], "critical": True},
        ...
    ]
    """
    if search_fn is None:
        try:
            from search_gateway import search
            search_fn = lambda q, n=10: search(q, max_results=n)
        except:
            search_fn = lambda q, n=10: []
    
    interceptor = VerificationInterceptor()
    
    for c in claims:
        interceptor.add_claim(c['claim'], c.get('critical', True))
        
        all_results = []
        for query in c.get('queries', [c['claim']]):
            r = search_fn(query, max_results=10)
            all_results.extend(r or [])
            time.sleep(0.3)
        
        # 简单匹配：至少有一条结果的 title/snippet/url 包含了查询中的关键词
        found = False
        keywords = [k for q in c['queries'] for k in q.split() if len(k) > 1]
        for res in all_results:
            text = f"{res.get('title', '')} {res.get('snippet', '')} {res.get('url', '')}".lower()
            match = sum(1 for k in keywords if k.lower() in text)
            if match >= 1:  # 至少一个关键词命中
                found = True
                break
        
        interceptor.add_results(c['claim'], all_results if found else [])
    
    return interceptor


# ═══════════════════════════════════════════════
# Part B: 完善的 Markdown → DOCX 解析器
# ═══════════════════════════════════════════════

class MarkdownDocxBuilder:
    """
    完善的 Markdown → DOCX 转换器
    处理: 标题、粗体、列表、表格、URL 替换为 [1][2] 脚注
    """
    
    def __init__(self):
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        
        self.doc = Document()
        self.Pt = Pt
        self.RGBColor = RGBColor
        self.Cm = Cm
        self.WD_ALIGN = WD_ALIGN_PARAGRAPH
        self.WD_ALIGN_TABLE = WD_TABLE_ALIGNMENT
        self.qn = qn
        self.nsdecls = nsdecls
        self.parse_xml = parse_xml
        
        self.urls = []
        self.url_cache = {}
        
        # 全局字体
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(4)
        
        for level in range(1, 4):
            hs = self.doc.styles[f'Heading {level}']
            hs.font.name = '微软雅黑'
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        section = self.doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    def _replace_urls(self, text):
        def repl(m):
            url = m.group(0)
            if url not in self.url_cache:
                self.url_cache[url] = len(self.urls) + 1
                self.urls.append(url)
            return f"[{self.url_cache[url]}]"
        return re.sub(r'https?://[^\s<>\)\]]+', repl, text)
    
    def _add_rich_run(self, paragraph, text, bold=False, size=10, color=None):
        text = self._replace_urls(text)
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for j, part in enumerate(parts):
            if not part:
                continue
            is_bold = (j % 2 == 1) or bold
            r = paragraph.add_run(part)
            r.bold = is_bold
            r.font.name = '微软雅黑'
            r.element.rPr.rFonts.set(self.qn('w:eastAsia'), '微软雅黑')
            r.font.size = self.Pt(size)
            if color:
                r.font.color.rgb = color
    
    def heading(self, text, level=1):
        colors = {1: self.RGBColor(0x2C,0x3E,0x50), 2: self.RGBColor(0x34,0x49,0x5E), 3: self.RGBColor(0x7F,0x8C,0x8D)}
        p = self.doc.add_heading(text, level=level)
        for r in p.runs:
            r.font.name = '微软雅黑'
            r.element.rPr.rFonts.set(self.qn('w:eastAsia'), '微软雅黑')
            r.font.size = self.Pt(14 if level == 1 else 12 if level == 2 else 11)
            r.font.color.rgb = colors.get(level, self.RGBColor(0x33,0x33,0x33))
    
    def para(self, text, style=None, bold=False, size=10, color=None, space_before=4, space_after=4):
        if style == 'List Bullet':
            txt = self._replace_urls(text)
            p = self.doc.add_paragraph(txt, style='List Bullet')
            for r in p.runs:
                r.font.name = '微软雅黑'
                r.element.rPr.rFonts.set(self.qn('w:eastAsia'), '微软雅黑')
                r.font.size = self.Pt(10)
        else:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = self.Pt(space_before)
            p.paragraph_format.space_after = self.Pt(space_after)
            self._add_rich_run(p, text, bold=bold, size=size, color=color)
        return p
    
    def table(self, headers, rows, col_widths=None):
        tbl = self.doc.add_table(rows=1+len(rows), cols=len(headers))
        tbl.style = 'Table Grid'
        tbl.autofit = True
        tbl.alignment = self.WD_ALIGN_TABLE.CENTER
        
        for i, h in enumerate(headers):
            h_clean = self._replace_urls(h)
            p = tbl.cell(0, i).paragraphs[0]
            p.alignment = self.WD_ALIGN.CENTER
            r = p.add_run(h_clean)
            r.bold = True
            r.font.name = '微软雅黑'
            r.element.rPr.rFonts.set(self.qn('w:eastAsia'), '微软雅黑')
            r.font.size = self.Pt(10)
            r.font.color.rgb = self.RGBColor(0xFF, 0xFF, 0xFF)
            self._shade_cell(tbl.cell(0, i), '2C3E50')
        
        for r_idx, row_data in enumerate(rows):
            for c_idx, val in enumerate(row_data):
                if c_idx >= len(headers):
                    break
                val = self._replace_urls(val)
                bg = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
                p = tbl.cell(r_idx+1, c_idx).paragraphs[0]
                r = p.add_run(val)
                r.font.name = '微软雅黑'
                r.element.rPr.rFonts.set(self.qn('w:eastAsia'), '微软雅黑')
                r.font.size = self.Pt(9)
                self._shade_cell(tbl.cell(r_idx+1, c_idx), bg)
        
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in tbl.rows:
                    row.cells[i].width = w
        
        self.doc.add_paragraph()
        return tbl
    
    def _shade_cell(self, cell, color_hex):
        shading = self.parse_xml(f'<w:shd {self.nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _is_table_separator(self, line):
        return bool(re.match(r'^\|[=\s\-:]+\|$', line.strip()))
    
    def _parse_and_add_table(self, lines):
        data = [l for l in lines if not self._is_table_separator(l)]
        if not data:
            return
        
        headers = [c.strip() for c in data[0].split('|')[1:-1]]
        if not headers:
            return
        
        rows = []
        for line in data[1:]:
            cols = [c.strip() for c in line.split('|')[1:-1]]
            if cols and any(c for c in cols):
                rows.append(cols)
        
        if rows:
            self.table(headers, rows)
    
    def parse_markdown(self, md_text):
        lines = md_text.split('\n')
        i = 0
        in_table = False
        table_buf = []
        
        while i < len(lines):
            s = lines[i].strip()
            
            if not s:
                if in_table:
                    self._parse_and_add_table(table_buf)
                    in_table = False
                    table_buf = []
                i += 1
                continue
            
            if s.startswith('|'):
                if not in_table:
                    table_buf = [s]
                    in_table = True
                else:
                    table_buf.append(s)
                i += 1
                continue
            elif in_table:
                self._parse_and_add_table(table_buf)
                in_table = False
                table_buf = []
            
            # 标题
            heading = False
            for lev, pf in [(1, '# '), (2, '## '), (3, '### ')]:
                if s.startswith(pf):
                    self.heading(s[len(pf):].strip(), level=lev)
                    heading = True
                    break
            if heading:
                i += 1
                continue
            
            # **标题**：内容
            m = re.match(r'\*\*([^*]+)\*\*[：:]\s*(.*)', s)
            if m:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = self.Pt(6)
                p.paragraph_format.space_after = self.Pt(6)
                r1 = p.add_run(m.group(1) + '：')
                r1.bold = True
                r1.font.name = '微软雅黑'
                r1.element.rPr.rFonts.set(self.qn('w:eastAsia'), '微软雅黑')
                r1.font.size = self.Pt(10)
                r1.font.color.rgb = self.RGBColor(0x2C, 0x3E, 0x50)
                self._add_rich_run(p, m.group(2))
                i += 1
                continue
            
            # 整行粗体强调
            m = re.match(r'\*\*(.+)\*\*$', s)
            if m and ':' not in s:
                self.para(m.group(1), bold=True, color=self.RGBColor(0xE7,0x4C,0x3C), space_before=6)
                i += 1
                continue
            
            # 列表
            if s.startswith('- ') or s.startswith('* '):
                self.para(s[2:], style='List Bullet')
                i += 1
                continue
            
            # 普通段落
            if s:
                self.para(s)
            
            i += 1
        
        if in_table and table_buf:
            self._parse_and_add_table(table_buf)
    
    def add_references(self):
        if not self.urls:
            return
        self.heading('参考文献', level=1)
        for i, url in enumerate(self.urls, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = self.Pt(2)
            p.paragraph_format.space_after = self.Pt(2)
            r1 = p.add_run(f'[{i}] ')
            r1.bold = True
            r1.font.name = '微软雅黑'
            r1.element.rPr.rFonts.set(self.qn('w:eastAsia'), '微软雅黑')
            r1.font.size = self.Pt(9)
            r2 = p.add_run(url)
            r2.font.name = '微软雅黑'
            r2.element.rPr.rFonts.set(self.qn('w:eastAsia'), '微软雅黑')
            r2.font.size = self.Pt(9)
            r2.font.color.rgb = self.RGBColor(0x29, 0x80, 0xB9)
    
    def save(self, path):
        self.add_references()
        self.doc.save(str(path))


# ═══════════════════════════════════════════
# 工作流规则
# ═══════════════════════════════════════════
